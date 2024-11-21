import os
from flask import Flask, request, render_template, send_file, redirect, url_for
from openpyxl import load_workbook
from docx import Document
import zipfile

# 初始化 Flask 应用
app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def replace_text_with_style(paragraph, replacements):
    """
    替换段落中的占位符，同时保留原有样式
    """
    # 将整个段落文本合并
    full_text = ''.join(run.text for run in paragraph.runs)

    # 替换占位符
    for key, value in replacements.items():
        placeholder = f"{{{key}}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))

    # 清空原有段落的内容
    for run in paragraph.runs:
        run.text = ""

    # 重新分配内容到第一个 Run，保持样式一致
    if paragraph.runs:
        paragraph.runs[0].text = full_text

def replace_table_text_with_style(table, replacements):
    """
    替换表格中的占位符，同时保留原有样式
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_with_style(paragraph, replacements)

# 替换字段的函数
def replace_fields_in_vertical_word(word_file, excel_file, output_folder):
    """
    替换竖排 Word 模板中的占位符并生成独立的文档
    """
    wb = load_workbook(excel_file)
    generated_files = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        headers = list(sheet.iter_rows(min_row=4, max_row=4, values_only=True))[0]
        data_rows = list(sheet.iter_rows(min_row=5, values_only=True))

        for i, row in enumerate(data_rows, start=1):
            replacements = dict(zip(headers, row))

            if not replacements.get("Name") or not replacements.get("Address"):
                continue

            doc = Document(word_file)

            for paragraph in doc.paragraphs:
                replace_text_with_style(paragraph, replacements)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_with_style(paragraph, replacements)

            output_file = os.path.join(output_folder, f"{sheet_name}_output_{i}.docx")
            doc.save(output_file)
            generated_files.append(output_file)

    return generated_files

# 首页路由
@app.route("/", methods=["GET", "POST"])
def upload_files():
    if request.method == "POST":
        word_file = request.files["word_file"]
        excel_file = request.files["excel_file"]

        if word_file and excel_file:
            # 保存上传的文件
            word_path = os.path.join(UPLOAD_FOLDER, word_file.filename)
            excel_path = os.path.join(UPLOAD_FOLDER, excel_file.filename)
            word_file.save(word_path)
            excel_file.save(excel_path)

            # 生成多个文档
            generated_files = replace_fields_in_vertical_word(word_path, excel_path, OUTPUT_FOLDER)

            if not generated_files:
                return "未生成任何文档，请检查模板和数据是否匹配。", 400

            # 将生成的文件打包成 ZIP 返回
            zip_filename = "output_documents.zip"
            zip_path = os.path.join(OUTPUT_FOLDER, zip_filename)
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file in generated_files:
                    if os.path.exists(file):
                        zipf.write(file, os.path.basename(file))
                        print(f"打包文件：{file}")

            return redirect(url_for("download_file", filename=zip_filename))
    return render_template("upload.html")

# 文件下载路由
@app.route("/download/<filename>")
def download_file(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        return "文件未找到。", 404
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)